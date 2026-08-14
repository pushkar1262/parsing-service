"""DOCX, where heading structure is stated rather than guessed.

This is the format worth getting exactly right first. A .docx carries paragraph
*style names* — "Heading 1", "List Bullet", "Title" — so the section tree is
recoverable with certainty rather than inferred from font sizes the way it has to be
in a PDF. Whatever this parser gets wrong is a bug; whatever the PDF parser gets
wrong about headings is, at best, a heuristic that lost.

Two things python-docx does not give us and we have to do ourselves:

`document.paragraphs` and `document.tables` are separate collections, so reading
them in turn puts every table at the end of the document. Tables in a requirements
doc are usually *the* requirements, and a table detached from the heading that
introduced it loses the context that made it meaningful. So we walk the body XML in
document order instead.

`w:numPr` gives a list level but the numbering definitions that say whether the
level is bulleted or numbered live in a separate part. We read the level (which is
what nesting needs) and infer ordered-ness from the style name, which is right for
documents produced by Word's own list styles and degrades to "bulleted" otherwise —
a wrong bullet is cosmetic, a wrong nesting level corrupts the structure.
"""

from __future__ import annotations

import io
from collections.abc import Iterator
from typing import Any

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from domain.document import BlockType, TableData
from parse.base import CorruptDocument, ParseResult, RawBlock
from parse.normalise import clean_inline

# Word's built-in heading styles, plus the outline level they imply.
_HEADING_PREFIXES = ("heading", "titre", "berschrift")  # en / fr / de(umlaut-stripped)

_LIST_STYLE_HINTS = ("list bullet", "list number", "list paragraph")


def _iter_body(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """Paragraphs and tables in the order they appear in the document.

    The reason this function exists rather than `document.paragraphs + tables`: the
    latter reorders content, and a table's meaning usually depends on the heading
    immediately above it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _style_name(paragraph: Paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip().lower()
    except (AttributeError, KeyError):  # a style referencing a missing definition
        return ""


def _heading_level(style: str) -> int | None:
    """The outline level a style name implies, or None if it is not a heading.

    Matches on prefix rather than exact name because real documents carry
    "Heading 2", "Heading 2 Char", and localised variants, all of which mean the
    same thing structurally.
    """
    if style in ("title", "subtitle"):
        return 1 if style == "title" else 2
    for prefix in _HEADING_PREFIXES:
        if style.startswith(prefix):
            digits = "".join(c for c in style if c.isdigit())
            if digits:
                return min(int(digits[0]), 6)
    return None


def _list_level(paragraph: Paragraph) -> int | None:
    """The zero-based indent level from `w:numPr`, or None if not a list item.

    `w:ilvl` is the only reliable nesting signal in a .docx: the visual indent is a
    style property and can be anything, but the level is what Word itself uses to
    decide numbering.
    """
    pr = paragraph._p.pPr
    if pr is None:
        return None
    num_pr = pr.numPr
    if num_pr is None:
        return None
    ilvl = num_pr.ilvl
    if ilvl is None:
        return 0
    try:
        return max(int(ilvl.val), 0)
    except (TypeError, ValueError):
        return 0


def _style_list_level(style: str) -> int | None:
    """Nesting level from a list *style name*, for documents with no direct `w:numPr`.

    Both paths are needed and neither is redundant. Word's own UI attaches numbering
    directly to the paragraph, so `w:numPr` is present and authoritative. But
    documents generated from a template — including everything python-docx produces
    with `style="List Bullet 2"` — carry numbering on the style definition instead,
    leaving the paragraph bare. Reading only `w:numPr` flattens every level of those
    documents into one, which is the structural loss the block tree exists to avoid.

    "List Bullet" is level 0, "List Bullet 2" is level 1: Word numbers its style
    variants from 1 while `w:ilvl` counts from 0.
    """
    if not any(hint in style for hint in _LIST_STYLE_HINTS):
        return None
    digits = "".join(c for c in style if c.isdigit())
    if not digits:
        # "List Paragraph" with no numbering is Word's "indented text"; treat it as a
        # top-level item rather than dropping the content.
        return 0
    return max(int(digits) - 1, 0)


def _cell_text(cell: Any) -> str:
    """A cell flattened to one line.

    A cell containing several paragraphs becomes one space-joined line, because a
    pipe table row cannot hold a line break and the alternative — dropping all but
    the first paragraph — silently loses content.
    """
    return clean_inline(" ".join(p.text for p in cell.paragraphs))


def _table_block(table: Table) -> RawBlock | None:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([_cell_text(cell) for cell in row.cells])
    if not any(any(cell for cell in row) for row in rows):
        return None
    return RawBlock(type=BlockType.TABLE, table=TableData(rows=rows, header_rows=1))


class DocxParser:
    name = "docx"
    version = "1.0"
    media_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def parse(self, data: bytes, *, filename: str | None = None) -> ParseResult:
        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:  # python-docx raises a variety of types here
            raise CorruptDocument(f"could not open as DOCX: {exc}") from exc

        result = ParseResult()
        root: list[RawBlock] = []
        heading_stack: list[tuple[int, RawBlock]] = []
        # One open list per nesting level, so consecutive items join a single list
        # rather than each becoming a list of one.
        list_stack: list[tuple[int, RawBlock]] = []

        def sink() -> list[RawBlock]:
            return heading_stack[-1][1].children if heading_stack else root

        def close_lists() -> None:
            list_stack.clear()

        for item in _iter_body(document):
            if isinstance(item, Table):
                close_lists()
                block = _table_block(item)
                if block is not None:
                    sink().append(block)
                continue

            text = clean_inline(item.text)
            style = _style_name(item)

            if not text:
                continue

            level = _heading_level(style)
            if level is not None:
                close_lists()
                heading = RawBlock(type=BlockType.HEADING, text=text, depth=level)
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                sink().append(heading)
                heading_stack.append((level, heading))
                continue

            list_level = _list_level(item)
            if list_level is None:
                list_level = _style_list_level(style)

            if list_level is not None:
                self._append_list_item(
                    text, list_level, style, list_stack, sink()
                )
                continue

            close_lists()
            sink().append(RawBlock(type=BlockType.PARAGRAPH, text=text))

        result.blocks = root
        result.metadata.update(self._metadata(document, root))
        return result

    # ------------------------------------------------------------------ lists

    def _append_list_item(
        self,
        text: str,
        level: int,
        style: str,
        list_stack: list[tuple[int, RawBlock]],
        sink: list[RawBlock],
    ) -> None:
        ordered = "number" in style
        while list_stack and list_stack[-1][0] > level:
            list_stack.pop()

        if not list_stack:
            container = RawBlock(type=BlockType.LIST, attrs={"ordered": ordered})
            sink.append(container)
            list_stack.append((level, container))
        elif list_stack[-1][0] < level:
            parent_items = list_stack[-1][1].children
            container = RawBlock(type=BlockType.LIST, attrs={"ordered": ordered})
            if parent_items:
                parent_items[-1].add(container)
            else:
                list_stack[-1][1].add(container)
            list_stack.append((level, container))

        container = list_stack[-1][1]
        depth = len(list_stack)
        attrs: dict[str, Any] = {"ordered": ordered}
        if ordered:
            attrs["number"] = sum(
                1 for child in container.children if child.type is BlockType.LIST_ITEM
            ) + 1
        container.add(
            RawBlock(type=BlockType.LIST_ITEM, text=text, depth=depth, attrs=attrs)
        )

    # --------------------------------------------------------------- metadata

    def _metadata(self, document: DocxDocument, root: list[RawBlock]) -> dict[str, Any]:
        """Core properties, with the empty strings Word writes treated as absent.

        Word populates these fields with `""` rather than omitting them, so a naive
        read reports every document as having an author named nothing. `page_count`
        is deliberately never set: it cannot be known without rendering the document,
        and a guess here would be worse than the honest `None`.
        """
        core = document.core_properties
        meta: dict[str, Any] = {"format": "docx"}

        def put(key: str, value: Any) -> None:
            if isinstance(value, str):
                value = value.strip()
            if value:
                meta[key] = value

        put("title", core.title)
        put("author", core.author)
        put("subject", core.subject)
        put("language", core.language)
        put("created_at", core.created)
        put("modified_at", core.modified)
        if core.keywords:
            keywords = [k.strip() for k in core.keywords.replace(";", ",").split(",")]
            put("keywords", [k for k in keywords if k])

        if "title" not in meta:
            for block in root:
                if block.type is BlockType.HEADING:
                    put("title", block.text)
                break
        return meta
